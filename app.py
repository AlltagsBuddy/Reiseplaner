# reiseplaner/app.py

from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from openai import OpenAI
import os
import logging

from dotenv import load_dotenv
load_dotenv()


# === Konfiguration ===

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__)
CORS(app)

logging.basicConfig(level=logging.INFO)

# === Routen ===

@app.route('/')
def index():
    return render_template('reiseplaner.html')


@app.route('/api/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()
        inhalt = data.get('inhalt', '').strip()
        stil = data.get('stil', 'Strand').strip()

        if not inhalt:
            return jsonify({'error': 'Keine Inhalte übergeben'}), 400

        prompt = f"""
Du bist ein smarter, professioneller Reiseassistent. Der Nutzer plant eine Reise im Stil: {stil}.

Analysiere die Angaben des Nutzers sorgfältig und erstelle basierend darauf einen umfassenden, realistischen und alltagstauglichen Reiseplan.

Inhalt des Plans:

1. **Zielregion & Kurzbeschreibung**  
   Gib ein konkretes Reiseziel an, das zu den Angaben passt, inkl. kurzer Beschreibung der Region und Besonderheiten.

2. **Aktivitäten & Highlights**  
   Liste 3–5 passende Aktivitäten oder Sehenswürdigkeiten auf, die zum Stil und Reiseziel passen.

3. **Empfohlener Reisezeitraum**  
   Gib den optimalen Zeitraum (Monate, Wetterlage, Touristenaufkommen) an.

4. **Reisebudget (realistisch)**  
   Schätze die Gesamtkosten (Hin- & Rückreise, Unterkunft, Verpflegung, Aktivitäten, ggf. Maut und Eintritte). Gib auch an, was ggf. nicht im Budget enthalten ist.

5. **Detaillierte Reiseroute**  
   Gib eine logische, mehrtägige Route an (mit Zwischenstopps) inkl. sinnvoller Reihenfolge der Orte. Berücksichtige Mautpflicht, Fahrzeiten und Highlights entlang der Strecke.

6. **Tagesablauf (Beispiel)**  
   Erstelle beispielhaft einen typischen Tagesablauf (z. B. Frühstück – Aktivität – Mittag – Entspannung – Abendprogramm) für 1–2 Reisetage.

7. **Zusätzliche Tipps & Hinweise**  
   Was sollte man mitnehmen (z. B. Wanderschuhe, Adapter, Mückenspray)? Gibt es Impfempfehlungen, Visumspflicht, Versicherungen oder spezielle kulturelle Gepflogenheiten? Erwähne auch mögliche Extrakosten (Parken, Trinkgeld, Touristensteuer etc.).

8. **Notfall-Infos vor Ort**  
   Liste wichtige Notfallnummern (z. B. Polizei, Notruf, deutsche Botschaft), gängige Zahlungsmittel und ob Kreditkarten akzeptiert werden.

Angaben des Nutzers:
{inhalt}

Antworte im Klartext, gut lesbar, strukturiert mit Zwischenüberschriften. Kein Fließtext, sondern logisch gegliederte Abschnitte wie in einem echten Reiseplan.
"""

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Du bist ein professioneller Reiseberater und planst perfekte Urlaube."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        plan = response.choices[0].message.content.strip()
        return jsonify({'reiseplan': plan})

    except Exception as e:
        logging.exception("Fehler bei der Reiseplan-Generierung")
        return jsonify({'error': str(e)}), 500


@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    try:
        data = request.get_json()
        plan = data.get('reiseplan', 'Kein Inhalt')

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        doc.add_heading('Dein Reiseplan', level=1)
        for line in plan.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="reiseplan.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logging.exception("Fehler beim DOCX-Export")
        return jsonify({'error': str(e)}), 500


@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    try:
        data = request.get_json()
        plan = data.get('reiseplan', 'Kein Inhalt')

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm
        )

        styles = getSampleStyleSheet()
        story = [Paragraph("Dein Reiseplan", styles['Heading1'])]

        for line in plan.split('\n'):
            if line.strip():
                story.append(Paragraph(line.strip(), styles['Normal']))
                story.append(Spacer(1, 6))

        doc.build(story)

        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="reiseplan.pdf",
            mimetype="application/pdf"
        )

    except Exception as e:
        logging.exception("Fehler beim PDF-Export")
        return jsonify({'error': str(e)}), 500


# === Startserver ===
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get("PORT", 10000)))
